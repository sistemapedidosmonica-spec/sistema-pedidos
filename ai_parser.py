import anthropic
import json
import re
import unicodedata
from typing import Optional
import pdfplumber
import io


def _norm(texto: str) -> str:
    """Remove acentos e converte para minúsculas para comparação robusta."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', texto)
        if unicodedata.category(c) != 'Mn'
    ).lower()

# -------------------------------------------------------
# Filtro de pedidos de PREFEITURA / ÓRGÃO PÚBLICO
# O critério é: quem pediu (prefeitura) + é um pedido?
# Não importa o produto — pode ser alimento, uniforme,
# limpeza, carne, gênero, etc.
# -------------------------------------------------------

# Indicadores de que é um PEDIDO / ORDEM DE COMPRA
_PALAVRAS_PEDIDO = [
    'pedido', 'ordem de compra', 'requisição', 'requisicao',
    'solicitação', 'solicitacao', 'nota de empenho', 'empenho',
    'fornecimento', 'aquisição', 'aquisicao', 'compra',
    'lista de', 'relação de', 'relacao de',
    'nota de pedido', 'oferta de compra',
    'pregão', 'pregao', 'licitação', 'licitacao',
    'cotação', 'cotacao', 'proposta',
]

# Indicadores de que o remetente é uma PREFEITURA / ÓRGÃO PÚBLICO
_PALAVRAS_ORGAO = [
    'prefeitura', 'pref.', 'municipio', 'município',
    'secretaria', 'setor de compras', 'setor de licitação',
    'departamento de compras', 'dept. compras',
    'câmara municipal', 'camara municipal',
    'governo municipal', 'administração municipal',
    'fundo municipal', 'autarquia',
]

# Descarte imediato — spam, marketing, sistemas automáticos e processos de licitação
_PALAVRAS_DESCARTE = [
    # Spam / marketing
    'unsubscribe', 'cancelar inscrição', 'cancelar inscricao',
    'newsletter', 'promoção especial', 'oferta especial',
    'black friday', 'cyber monday',
    'fatura do cartão', 'vencimento da fatura', 'boleto vencendo',
    'confirmação de cadastro', 'bem-vindo ao', 'obrigado por se cadastrar',
    'rastreamento do pedido', 'código de rastreio', 'entrega realizada',
    'noreply@', 'no-reply@', 'donotreply@', 'mailer-daemon',
    'sua assinatura', 'renovação automática',
    'sebrae', 'senai', 'sesi', 'fiemg', 'fecomércio',
    # Processos de licitação / contratos que NÃO são pedidos de compra
    'aditivar', 'aditivo', 'avaliação de contrato', 'avaliacao de contrato',
    'ata de registro', 'ata de licitação', 'ata de licitacao',
    'arp ', '/arp', 'resultado de licitação', 'resultado de licitacao',
    'edital', 'impugnação', 'impugnacao', 'recurso de licitação',
    'habilitação', 'habilitacao', 'inabilitação', 'inabilitacao',
    'homologação', 'homologacao', 'adjudicação', 'adjudicacao',
    'sessão pública', 'sessao publica', 'pregão eletrônico',
    'dispensa de licitação', 'dispensa de licitacao',
    'nota de esclarecimento', 'esclarecimento ao edital',
    'contrato administrativo', 'termo de contrato',
    'termo aditivo', 'rescisão contratual', 'rescisao contratual',
]


def email_e_pedido(email_data: dict) -> tuple[bool, str]:
    """
    Verifica se um email é um pedido/ordem de compra de prefeitura ou órgão público.
    Retorna (True/False, motivo).

    Aceita qualquer tipo de produto: alimentos, carne, uniforme,
    limpeza, gêneros — o que importa é ser uma prefeitura pedindo.
    """
    assunto   = (email_data.get('assunto') or '').lower()
    corpo     = (email_data.get('texto_corpo') or '').lower()
    remetente = (email_data.get('email_remetente') or '').lower()
    texto     = assunto + ' ' + corpo[:5000]

    # 1. Descarte imediato
    for palavra in _PALAVRAS_DESCARTE:
        if palavra in texto or palavra in remetente:
            return False, f'spam/automático: "{palavra}"'

    pontos = 0
    motivos = []

    # 2. Domínio .gov.br = prefeitura confirmada (peso alto)
    if '.gov.br' in remetente:
        pontos += 4
        motivos.append('domínio .gov.br')

    # 3. Palavra de órgão público no remetente ou assunto
    for p in _PALAVRAS_ORGAO:
        if p in remetente or p in assunto:
            pontos += 3
            motivos.append(f'órgão: "{p}"')
            break

    # 4. Palavra de pedido/ordem de compra no assunto ou corpo
    for p in _PALAVRAS_PEDIDO:
        if p in texto:
            pontos += 2
            motivos.append(f'pedido: "{p}"')
            break

    # 5. Tem anexo formal (PDF, Word, Excel, CSV)
    for anexo in email_data.get('anexos', []):
        nome = (anexo.get('nome') or '').lower()
        tipo = (anexo.get('tipo') or '').lower()
        if any(ext in nome for ext in ['.pdf', '.doc', '.xls', '.xlsx', '.csv']):
            pontos += 2
            motivos.append('anexo formal')
            break
        elif any(t in tipo for t in ['pdf', 'word', 'excel', 'sheet']):
            pontos += 2
            motivos.append('anexo formal')
            break

    # Aceita com pelo menos 3 pontos (evita falsos positivos)
    if pontos >= 3:
        return True, ' | '.join(motivos)

    return False, f'sem evidência de pedido de prefeitura (pontos: {pontos})'


def extrair_texto_pdf(dados_pdf: bytes) -> str:
    """Extrai texto de um PDF em bytes."""
    try:
        with pdfplumber.open(io.BytesIO(dados_pdf)) as pdf:
            texto = ''
            for pagina in pdf.pages:
                t = pagina.extract_text()
                if t:
                    texto += t + '\n'
        return texto.strip()
    except Exception as e:
        return f'[Erro ao ler PDF: {e}]'


def extrair_texto_docx(dados_docx: bytes) -> str:
    """Extrai texto de um arquivo Word (.docx) em bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(dados_docx))
        linhas = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tenta também tabelas
        for table in doc.tables:
            for row in table.rows:
                células = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if células:
                    linhas.append(' | '.join(células))
        return '\n'.join(linhas)
    except Exception as e:
        return f'[Erro ao ler Word: {e}]'


def extrair_texto_excel(dados_excel: bytes) -> str:
    """Extrai texto de um arquivo Excel (.xls ou .xlsx) em bytes."""
    try:
        import openpyxl
        import io
        wb = openpyxl.load_workbook(io.BytesIO(dados_excel), read_only=True, data_only=True)
        linhas = []
        for sheet in wb.worksheets:
            linhas.append(f'[Planilha: {sheet.title}]')
            for row in sheet.iter_rows(values_only=True):
                células = [str(c) for c in row if c is not None and str(c).strip()]
                if células:
                    linhas.append(' | '.join(células))
        return '\n'.join(linhas)
    except Exception:
        try:
            import xlrd, io
            wb = xlrd.open_workbook(file_contents=dados_excel)
            linhas = []
            for sheet in wb.sheets():
                linhas.append(f'[Planilha: {sheet.name}]')
                for r in range(sheet.nrows):
                    row = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)
                           if str(sheet.cell_value(r, c)).strip()]
                    if row:
                        linhas.append(' | '.join(row))
            return '\n'.join(linhas)
        except Exception as e:
            return f'[Erro ao ler Excel: {e}]'


def preparar_conteudo_email(email_data: dict) -> str:
    """
    Prepara o conteúdo completo de um email para enviar ao Claude.

    Suporta TODOS os formatos de pedido:
    - Texto direto no corpo do email
    - PDF anexado
    - Word (.doc, .docx) anexado
    - Excel (.xls, .xlsx) anexado
    - CSV ou TXT anexado
    - Qualquer combinação dos acima
    """
    conteudo = f"Assunto: {email_data.get('assunto', '')}\n"
    conteudo += f"Remetente: {email_data.get('email_remetente', '')}\n"
    conteudo += f"Data: {email_data.get('data_recebimento', '')}\n\n"

    # Corpo do email (texto direto ou HTML)
    corpo = email_data.get('texto_corpo', '') or ''
    if corpo.strip():
        conteudo += "=== CORPO DO EMAIL ===\n"
        conteudo += corpo + '\n'

    # Processa TODOS os anexos
    for anexo in email_data.get('anexos', []):
        nome = anexo.get('nome', '')
        nome_lower = nome.lower()
        tipo = (anexo.get('tipo', '') or '').lower()
        dados = anexo.get('dados', b'') or b''

        if not dados:
            continue

        conteudo += f"\n=== ANEXO: {nome} ===\n"

        if 'pdf' in tipo or nome_lower.endswith('.pdf'):
            conteudo += extrair_texto_pdf(dados)

        elif 'word' in tipo or nome_lower.endswith('.docx') or nome_lower.endswith('.doc'):
            conteudo += extrair_texto_docx(dados)

        elif any(x in tipo for x in ['excel', 'sheet']) or \
             nome_lower.endswith('.xlsx') or nome_lower.endswith('.xls'):
            conteudo += extrair_texto_excel(dados)

        elif 'text' in tipo or nome_lower.endswith('.csv') or nome_lower.endswith('.txt'):
            try:
                conteudo += dados.decode('utf-8', errors='replace')
            except Exception:
                conteudo += '[Não foi possível ler o arquivo de texto]'

        else:
            # Tenta ler como texto de qualquer forma
            try:
                texto = dados.decode('utf-8', errors='replace')
                if texto.strip():
                    conteudo += texto
                else:
                    conteudo += f'[Arquivo binário: {nome}]'
            except Exception:
                conteudo += f'[Arquivo não lido: {nome}]'

    return conteudo[:20000]  # Limita para não estourar o contexto


def extrair_empresa_prefeitura_do_assunto(assunto: str, prefeituras: list, empresas: list) -> tuple:
    """
    Extrai empresa e prefeitura diretamente do assunto do email.
    A funcionária coloca no assunto: "Nome da Cidade - Nome da Empresa"
    Ex: "Casa Grande - Profeta", "Itaverá - Eros", "Caranaíba - Eros"
    Retorna (prefeitura_obj, empresa_obj) ou (None, None).

    Usa normalização de acentos para comparação robusta:
    "São Brás" == "Sao Bras", "Itaverá" == "Itavera", etc.
    """
    if not assunto:
        return None, None

    assunto_norm = _norm(assunto)

    # Palavras que devem ser ignoradas no matching de prefeitura
    _STOPWORDS = {'prefeitura', 'municipal', 'secretaria', 'municipio', 'municipais'}

    # Encontra empresa pelo assunto (normalizado)
    empresa_encontrada = None
    for e in empresas:
        # Usa palavras significativas do nome da empresa (len >= 3 para pegar "Eros")
        palavras = [_norm(w) for w in e.nome.split() if len(w) >= 3]
        if any(p in assunto_norm for p in palavras):
            empresa_encontrada = e
            break

    # Encontra prefeitura pelo assunto (normalizado)
    prefeitura_encontrada = None
    melhor_score = 0
    for p in prefeituras:
        # Pega as palavras significativas do município (len >= 3, sem stopwords)
        palavras_mun = [
            _norm(w) for w in p.nome.split()
            if len(w) >= 3 and _norm(w) not in _STOPWORDS
        ]
        score = sum(1 for w in palavras_mun if w in assunto_norm)
        if score > melhor_score:
            melhor_score = score
            prefeitura_encontrada = p

    if melhor_score == 0:
        prefeitura_encontrada = None

    return prefeitura_encontrada, empresa_encontrada


def extrair_pedido_com_ia(email_data: dict, api_key: str) -> dict:
    """
    Usa Claude IA para extrair os itens do pedido de um email,
    independente do formato em que chegou.
    """
    if not api_key:
        return {'erro': 'Chave da API não configurada', 'itens': []}

    client = anthropic.Anthropic(api_key=api_key)
    conteudo = preparar_conteudo_email(email_data)

    prompt = f"""Você é um assistente especializado em processar pedidos de compra de prefeituras municipais brasileiras para distribuidoras de alimentos e outros produtos.

CONTEXTO IMPORTANTE:
Este email foi enviado por uma funcionária que consolida todos os pedidos semanais e os encaminha para o sistema.
Cada email pode conter o pedido de UMA prefeitura específica destinado a UMA das seguintes empresas distribuidoras:
- Profeta Distribuidora e Serviços
- Eros Distribuidora
- Fortumel Produtos

Prefeituras atendidas:
- Casa Grande → contrato com Profeta
- Queluzito → contrato com Profeta
- Resende Costa → contrato com Profeta
- Senhora dos Remédios → contrato com Profeta
- Entre Rios de Minas → contrato com Profeta
- São Brás do Suaçuí → contrato com Profeta
- Itaverá → contratos com Profeta E com Eros (verificar qual empresa está mencionada)
- Caranaíba → contrato com Eros

Analise o email abaixo e extraia todas as informações em JSON.

EMAIL:
{conteudo}

Retorne APENAS um JSON válido (sem texto antes ou depois):
{{
  "empresa_nome": "nome EXATO de uma das 3 empresas distribuidoras: 'Profeta Distribuidora e Servicos', 'Eros Distribuidora' ou 'Fortumel Produtos'. Use a tabela acima e o conteúdo do email para determinar. Se Itavera, verifique qual empresa está mencionada.",
  "prefeitura_nome": "nome completo da prefeitura/município (ex: Prefeitura Municipal de Casa Grande, Prefeitura Municipal de Queluzito)",
  "prefeitura_endereco": "endereço de entrega se mencionado, senão null",
  "prefeitura_contato": "nome do responsável se mencionado, senão null",
  "data_entrega": "data de entrega se mencionada (formato dd/mm/aaaa), senão null",
  "observacoes": "observações gerais do pedido se houver, senão null",
  "itens": [
    {{
      "produto": "nome do produto normalizado com maiúscula (ex: Banana Nanica, Alface Crespa, Frango Inteiro)",
      "quantidade": 10.5,
      "unidade": "kg, g, un, cx, fd (fardo), sc (saco), lt (litro)",
      "observacao": "observação do item se houver, senão null"
    }}
  ]
}}

Regras:
- Normalize nomes de produtos: primeira letra maiúscula, sem abreviações desnecessárias
- Se não houver itens de pedido, retorne itens como lista vazia []
- Quantidade deve ser número com ponto decimal (ex: 10.5)
- Unidade padrão para frutas/legumes/verduras é "kg"
- empresa_nome NUNCA pode ser null — use a tabela de contratos acima para determinar"""

    try:
        message = client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        resposta = message.content[0].text.strip()

        # Remove markdown code blocks se houver
        resposta = re.sub(r'^```json\s*', '', resposta)
        resposta = re.sub(r'^```\s*', '', resposta)
        resposta = re.sub(r'\s*```$', '', resposta)

        resultado = json.loads(resposta)
        return resultado

    except json.JSONDecodeError as e:
        return {'erro': f'Resposta inválida da IA: {e}', 'itens': []}
    except anthropic.AuthenticationError:
        return {'erro': 'Chave da API inválida. Verifique nas configurações.', 'itens': []}
    except Exception as e:
        return {'erro': f'Erro ao processar com IA: {str(e)}', 'itens': []}


def gerar_mensagem_whatsapp(itens_lista: list, nome_empresa: str = 'Hortifruti') -> str:
    """
    Gera mensagem formatada para WhatsApp para enviar aos fornecedores
    pedindo cotação de preços.
    """
    hoje = __import__('datetime').date.today().strftime('%d/%m/%Y')

    msg = f"*Olá! Sou da {nome_empresa}.*\n"
    msg += f"Preciso de uma cotação de preços para hoje {hoje}:\n\n"

    for i, item in enumerate(itens_lista, 1):
        qtd = item.get('quantidade_total', 0)
        unidade = item.get('unidade', 'kg')
        produto = item.get('produto', '')
        # Formata quantidade sem casas decimais desnecessárias
        qtd_fmt = int(qtd) if qtd == int(qtd) else f'{qtd:.1f}'
        msg += f"{i}. *{produto}* - {qtd_fmt} {unidade}\n"

    msg += "\nPor favor, me informe o preço por unidade de cada item. Obrigada!"
    return msg


def sugerir_categoria_produto(nome_produto: str) -> str:
    """
    Sugere a categoria de um produto para filtrar por fornecedor.
    Categorias: genero, carne, limpeza, hortifruti, uniforme
    """
    p = nome_produto.lower()

    # Carnes
    if any(kw in p for kw in [
        'frango', 'carne', 'boi', 'suino', 'suíno', 'porco', 'peixe',
        'filé', 'file', 'alcatra', 'peito', 'coxa', 'sobrecoxa', 'costela',
        'linguiça', 'linguica', 'salsicha', 'calabresa', 'patinho', 'músculo',
        'musculo', 'acém', 'acem', 'contrafilé', 'contrafile', 'pernil'
    ]):
        return 'carne'

    # Limpeza
    if any(kw in p for kw in [
        'detergente', 'sabão', 'sabao', 'desinfetante', 'alvejante',
        'limpador', 'sanitizante', 'multiuso', 'esponja', 'vassoura',
        'rodo', 'balde', 'pano', 'papel higienico', 'papel higiênico',
        'toalha papel', 'saco lixo', 'saco de lixo', 'cloro', 'hipoclorito',
        'sabonete', 'água sanitária', 'agua sanitaria', 'desincrustante',
        'flanela', 'lustra', 'cera', 'odorizante'
    ]):
        return 'limpeza'

    # Uniforme / EPI
    if any(kw in p for kw in [
        'uniforme', 'jaleco', 'avental', 'calçado', 'calcado', 'sapato',
        'bota', 'luva', 'capacete', 'epi', 'crachá', 'cracha',
        'camisa', 'calça', 'calca', 'bermuda', 'colete', 'touca'
    ]):
        return 'uniforme'

    # Hortifruti
    if any(kw in p for kw in [
        'banana', 'laranja', 'maçã', 'maca', 'manga', 'uva', 'melancia',
        'melão', 'melao', 'abacaxi', 'mamão', 'mamao', 'limão', 'limao',
        'tangerina', 'pêssego', 'pessego', 'ameixa', 'tomate', 'batata',
        'cebola', 'alho', 'cenoura', 'beterraba', 'abobrinha', 'pepino',
        'pimentão', 'pimentao', 'brócolis', 'brocolis', 'couve-flor',
        'couve flor', 'repolho', 'alface', 'rúcula', 'rucula', 'espinafre',
        'couve', 'agrião', 'agriao', 'vagem', 'milho', 'inhame', 'mandioca',
        'batata-doce', 'batata doce', 'jiló', 'jilo', 'quiabo', 'chuchu',
        'abóbora', 'abobora', 'maxixe', 'cará', 'cara'
    ]):
        return 'hortifruti'

    # Gênero (padrão: secos, laticínios, etc.)
    return 'genero'


def gerar_mensagem_fornecedor(itens_lista: list, fornecedor_nome: str,
                              categorias: list, nome_empresa: str = 'Hortifruti') -> str:
    """
    Gera mensagem WhatsApp para um fornecedor específico,
    filtrando apenas os itens das categorias que ele atende.
    """
    hoje = __import__('datetime').date.today().strftime('%d/%m/%Y')

    itens_filtrados = []
    for item in itens_lista:
        produto = item.get('produto', '')
        categoria = sugerir_categoria_produto(produto)
        if not categorias or categoria in categorias:
            itens_filtrados.append(item)

    if not itens_filtrados:
        return f'Nenhum produto desta lista é fornecido por {fornecedor_nome}.'

    msg = f"*Olá {fornecedor_nome}! Sou da {nome_empresa}.*\n"
    msg += f"Preciso de uma cotação de preços para hoje {hoje}:\n\n"

    for i, item in enumerate(itens_filtrados, 1):
        qtd = item.get('quantidade_total', 0)
        unidade = item.get('unidade', 'kg')
        produto = item.get('produto', '')
        qtd_fmt = int(qtd) if qtd == int(qtd) else f'{qtd:.1f}'
        msg += f"{i}. *{produto}* - {qtd_fmt} {unidade}\n"

    msg += "\nPor favor, me informe o preço por unidade de cada item. Obrigada!"
    return msg


def sugerir_destino_produto(nome_produto: str, api_key: str) -> str:
    """
    Sugere se um produto deve ser comprado no CEASA ou local.
    Usa regras simples para evitar chamadas desnecessárias à API.
    """
    produto_lower = nome_produto.lower()

    # Produtos tipicamente do CEASA (atacado especializado)
    ceasa_keywords = [
        'banana', 'laranja', 'maçã', 'manga', 'uva', 'melancia', 'melão',
        'abacaxi', 'mamão', 'limão', 'tangerina', 'pêssego', 'ameixa',
        'tomate', 'batata', 'cebola', 'alho', 'cenoura', 'beterraba',
        'abobrinha', 'pepino', 'pimentão', 'brócolis', 'couve-flor',
        'repolho', 'alface', 'rúcula', 'espinafre', 'couve', 'agrião',
        'vagem', 'ervilha', 'milho', 'inhame', 'cará', 'mandioca',
        'batata-doce', 'jiló', 'quiabo', 'maxixe', 'chuchu'
    ]

    # Produtos tipicamente locais (secos, grãos, laticínios)
    local_keywords = [
        'arroz', 'feijão', 'macarrão', 'farinha', 'açúcar', 'sal',
        'óleo', 'azeite', 'leite', 'queijo', 'manteiga', 'iogurte',
        'pão', 'biscoito', 'bolacha', 'café', 'achocolatado',
        'sardinha', 'atum', 'frango', 'carne', 'ovos'
    ]

    for kw in ceasa_keywords:
        if kw in produto_lower:
            return 'ceasa'

    for kw in local_keywords:
        if kw in produto_lower:
            return 'local'

    return 'local'  # padrão
